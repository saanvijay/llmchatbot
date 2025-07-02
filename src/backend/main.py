from dotenv import load_dotenv
from langchain_ollama import OllamaLLM
from langchain_core.prompts import ChatPromptTemplate
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from http import HTTPStatus
from functools import wraps
import logging
import os
from datetime import datetime
from collections import defaultdict
from langchain_chroma import Chroma
from langchain_ollama import OllamaEmbeddings
from langchain_core.documents import Document
import requests
import mimetypes
from werkzeug.utils import secure_filename
import tempfile
import PyPDF2
import docx
import speech_recognition as sr
import pyttsx3
import io
import base64
from pydub import AudioSegment
from gtts import gTTS

import os
import io
import pandas as pd
import chromadb


# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)
#global retriever 
#retriever = None

load_dotenv('config/.env')

# Configuration
app.config['OLLAMA_BASE_URL'] = os.getenv('OLLAMA_BASE_URL')
app.config['OLLAMA_MODEL'] = os.getenv('OLLAMA_MODEL')
app.config['OLLAMA_EMBEDDING_MODEL'] = os.getenv('OLLAMA_EMBEDDING_MODEL')
app.config['CONTEXT_EXPIRY'] = int(os.getenv('CONTEXT_EXPIRY'))
app.config['BACKEND_SERVER_PORT'] = int (os.getenv('BACKEND_SERVER_PORT'))
app.config['CHROMA_DB_PATH'] = os.getenv('CHROMA_DB_PATH')
app.config['CHROMA_COLLECTION'] = os.getenv('CHROMA_COLLECTION')
app.config['CHROMA_TELEMETRY_ANONYMIZED'] = os.getenv('CHROMA_TELEMETRY_ANONYMIZED')

# Initialize Ollama
template = """
Answer the queries based on the provided context.
Summary: {summary}
Question: {question}
Context: {context}
Answer: 
"""
model = OllamaLLM(base_url=app.config['OLLAMA_BASE_URL'], model=app.config['OLLAMA_MODEL'])
prompt = ChatPromptTemplate.from_template(template)
chain = prompt | model

# In-memory storage for context
context_store = defaultdict(lambda: {
    'context': '',
    'last_updated': datetime.now()
})

# In-memory storage for idempotency
idempotency_cache = {}

def validate_json(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not request.is_json:
            return jsonify({
                'status': 'error',
                'message': 'Content-Type must be application/json'
            }), HTTPStatus.BAD_REQUEST
        return f(*args, **kwargs)
    return decorated_function

def log_request(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        start_time = datetime.now()
        response = f(*args, **kwargs)
        duration = datetime.now() - start_time
        logger.info(f"Request: {request.method} {request.path} - Duration: {duration}")
        return response
    return decorated_function

def get_session_id():
    """Get or create a session ID from the request"""
    session_id = request.headers.get('X-Session-ID')
    if not session_id:
        session_id = request.remote_addr
    return session_id

def is_context_expired(session_id):
    """Check if the context for a session has expired"""
    if session_id not in context_store:
        return True
    
    last_updated = context_store[session_id]['last_updated']
    expiry_time = datetime.now() - last_updated
    return expiry_time.total_seconds() > app.config['CONTEXT_EXPIRY']

def get_idempotency_key():
    """Get idempotency key from request headers"""
    return request.headers.get('X-Idempotency-Key')

def is_idempotent_request(idempotency_key):
    """Check if this is a duplicate request based on idempotency key"""
    return idempotency_key in idempotency_cache

def cache_response(idempotency_key, response_data, status_code):
    """Cache the response for idempotency"""
    idempotency_cache[idempotency_key] = {
        'response': response_data,
        'status_code': status_code,
        'timestamp': datetime.now()
    }

def cleanup_old_cache():
    """Clean up old cached responses (older than 24 hours)"""
    current_time = datetime.now()
    keys_to_remove = []
    for key, value in idempotency_cache.items():
        if (current_time - value['timestamp']).total_seconds() > 86400:  # 24 hours
            keys_to_remove.append(key)
    
    for key in keys_to_remove:
        del idempotency_cache[key]

@app.route('/api/v1/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'success',
        'message': 'Service is healthy',
        'timestamp': datetime.now().isoformat()
    }), HTTPStatus.OK

@app.route('/api/v1/chat', methods=['POST'])
@validate_json
@log_request
def chat():
    try:
        # Check for idempotency
        idempotency_key = get_idempotency_key()
        if idempotency_key and is_idempotent_request(idempotency_key):
            # Return cached response
            cached_data = idempotency_cache[idempotency_key]
            return jsonify(cached_data['response']), cached_data['status_code']
        
        # Clean up old cache entries
        cleanup_old_cache()
        
        data = request.get_json()
        session_id = get_session_id()
        question = data.get("question", "")
        collection_name = data.get("collection_name")

        # Retrieve previous context or start fresh
        context = context_store[session_id]['context'] if session_id in context_store else ""
        summary = []
        if os.path.exists(app.config['CHROMA_DB_PATH']):
            client = chromadb.PersistentClient(path=app.config['CHROMA_DB_PATH'])
            embedding_fn = OllamaEmbeddings(base_url=app.config['OLLAMA_BASE_URL'], model=app.config['OLLAMA_EMBEDDING_MODEL'])
            chroma_collection = collection_name if collection_name else app.config['CHROMA_COLLECTION']
            vector_store = Chroma(
                client=client,
                collection_name=chroma_collection,
                embedding_function=embedding_fn
            )
            retriever = vector_store.as_retriever()
            #full_query = f"{context}\n\n{question}"
            summary = retriever.invoke(f"{question}")
        
        #logger.info(f"VIJAY:", {"summary": summary, "question": question, "context": context})
        if not summary and not context:
            chain1 = model
            result = chain1.invoke(question)
        elif not summary and context:
            chain2 = model
            full_query = f"{context}\n\n{question}"
            result = chain2.invoke(full_query)
        elif summary and not context:
            chain1 = model
            result = chain1.invoke(question)
        else:
            question1 = question
            summary1 = summary
            results = {}
            for word in question1.split():
                results[word] = any(word.lower() in str(w).lower() for w in summary1)
            if any(results.values()):
                result = chain.invoke({
                    "summary": str(summary), 
                    "question": question, 
                    "context": context
                })
            else:
                chain1 = model
                result = chain1.invoke(question)
        
        if result :
            #logger.info("VIJAY : Result is", {result})
            # Append new Q&A to context and store it
            context += f"Question: {question} Answer: {result}\n"

            context_store[session_id] = {
                'context': context,
                'last_updated': datetime.now()
            }
            
            response_data = {
                'status': 'success',
                'data': {
                    'question': question,
                    'answer': result,
                    'timestamp': datetime.now().isoformat(),
                    'session_id': session_id
                }
            }
            
            # Cache response if idempotency key is provided
            if idempotency_key:
                cache_response(idempotency_key, response_data, HTTPStatus.OK)
            
            return jsonify(response_data), HTTPStatus.OK

    except Exception as e:
        logger.error(f"Error processing chat request: {str(e)}")
        error_response = {
            'status': 'error',
            'message': 'Internal server error',
            'error': str(e)
        }
        
        # Cache error response if idempotency key is provided
        if idempotency_key:
            cache_response(idempotency_key, error_response, HTTPStatus.INTERNAL_SERVER_ERROR)
        
        return jsonify(error_response), HTTPStatus.INTERNAL_SERVER_ERROR

@app.route('/api/v1/context', methods=['DELETE'])
def clear_context():
    """Clear the context for a session"""
    session_id = get_session_id()
    if session_id in context_store:
        del context_store[session_id]
    return jsonify({
        'status': 'success',
        'message': 'Context cleared successfully'
    }), HTTPStatus.OK

@app.route('/api/v1/rag', methods=['POST'])
def rag():
    """Accepts a file (csv, pdf, doc) or a url, extracts content, and creates a Chroma DB collection for RAG."""
    try:
        collection_name = app.config['CHROMA_COLLECTION']
        if not collection_name:
            return jsonify({'status': 'error', 'message': 'CHROMA_COLLECTION is not configured in the environment.'}), HTTPStatus.INTERNAL_SERVER_ERROR

        # Check if it's a file upload
        if 'file' in request.files:
            file = request.files['file']
            filename = secure_filename(file.filename)
            ext = filename.split('.')[-1].lower()
            content = ''
            if ext == 'csv':
                df = pd.read_csv(file)
                # Concatenate all text columns for embedding
                content = '\n'.join(df.astype(str).apply(lambda row: ' '.join(row), axis=1))
            elif ext == 'pdf':
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                    file.save(tmp.name)
                    reader = PyPDF2.PdfReader(tmp.name)
                    content = '\n'.join(page.extract_text() or '' for page in reader.pages)
            elif ext in ['doc', 'docx']:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    file.save(tmp.name)
                    doc = docx.Document(tmp.name)
                    content = '\n'.join([para.text for para in doc.paragraphs])
            else:
                return jsonify({'status': 'error', 'message': 'Unsupported file type.'}), HTTPStatus.BAD_REQUEST
        else:
            # Assume JSON body for URL
            data = request.get_json()
            url = data.get('url')
            if not url:
                return jsonify({'status': 'error', 'message': 'url is required for RAG processing.'}), HTTPStatus.BAD_REQUEST
            headers = {
                    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.5 Safari/605.1.15"
            }
            response = requests.get(url, headers=headers)
            if response.status_code != 200:
                return jsonify({'status': 'error', 'message': f'Failed to fetch URL: {response.status_code}'}), HTTPStatus.BAD_REQUEST
            content = response.text

        if not content.strip():
            return jsonify({'status': 'error', 'message': 'No content extracted from input.'}), HTTPStatus.BAD_REQUEST

        # Create document and store in Chroma
        doc = Document(page_content=content, metadata={'source': collection_name})
        embeddings = OllamaEmbeddings(base_url=app.config['OLLAMA_BASE_URL'], model=app.config['OLLAMA_EMBEDDING_MODEL'])
        client = chromadb.PersistentClient(path=app.config['CHROMA_DB_PATH'])
        vectorstore = Chroma(
            client=client,
            collection_name=collection_name,
            embedding_function=embeddings
        )
        vectorstore.add_documents([doc], ids=['0'])
        return jsonify({'status': 'success', 'message': f'RAG collection {collection_name} created.'}), HTTPStatus.OK
    except Exception as e:
        logger.error(f"Error processing RAG request: {str(e)}")
        return jsonify({'status': 'error', 'message': 'Error processing input for RAG', 'error': str(e)}), HTTPStatus.INTERNAL_SERVER_ERROR

@app.route('/api/v1/voice/speech-to-text', methods=['POST'])
def speech_to_text():
    """Convert speech audio to text"""
    try:
        if 'audio' not in request.files:
            return jsonify({'status': 'error', 'message': 'No audio file provided'}), HTTPStatus.BAD_REQUEST
        
        audio_file = request.files['audio']
        if audio_file.filename == '':
            return jsonify({'status': 'error', 'message': 'No audio file selected'}), HTTPStatus.BAD_REQUEST
        
        # Save audio file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix='.audio') as tmp:
            audio_file.save(tmp.name)
            audio_path = tmp.name
        
        # Convert audio to WAV using pydub
        try:
            # Try to detect the format from the file
            audio = AudioSegment.from_file(audio_path)
            wav_path = audio_path.replace('.audio', '.wav')
            audio.export(wav_path, format="wav", parameters=["-ar", "16000", "-ac", "1"])
        except Exception as e:
            logger.error(f"Error converting audio format: {str(e)}")
            # Fallback: try to use the original file if conversion fails
            wav_path = audio_path
            logger.info("Using original audio file as fallback")
        
        # Initialize recognizer
        recognizer = sr.Recognizer()
        
        # Load audio file
        with sr.AudioFile(wav_path) as source:
            # Adjust for ambient noise
            recognizer.adjust_for_ambient_noise(source, duration=0.5)
            audio_data = recognizer.record(source)
        
        # Convert speech to text
        text = recognizer.recognize_google(audio_data)
        
        # Clean up temporary files
        os.unlink(audio_path)
        if wav_path != audio_path:  # Only delete wav file if it's different from original
            os.unlink(wav_path)
        
        return jsonify({
            'status': 'success',
            'text': text
        }), HTTPStatus.OK
        
    except sr.UnknownValueError:
        return jsonify({'status': 'error', 'message': 'Could not understand audio. Please speak clearly and try again.'}), HTTPStatus.BAD_REQUEST
    except sr.RequestError as e:
        return jsonify({'status': 'error', 'message': f'Speech recognition service error: {str(e)}'}), HTTPStatus.INTERNAL_SERVER_ERROR
    except Exception as e:
        logger.error(f"Error in speech-to-text: {str(e)}")
        return jsonify({'status': 'error', 'message': 'Error processing audio. Please try again.'}), HTTPStatus.INTERNAL_SERVER_ERROR

@app.route('/api/v1/voice/text-to-speech', methods=['POST'])
@validate_json
def text_to_speech():
    """Convert text to speech audio using gTTS (Google Text-to-Speech) and return MP3 audio."""
    try:
        data = request.get_json()
        text = data.get('text', '')
        
        if not text:
            return jsonify({'status': 'error', 'message': 'No text provided'}), HTTPStatus.BAD_REQUEST
        
        # Generate speech using gTTS
        tts = gTTS(text=text, lang='en', slow=False)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as tmp:
            audio_path = tmp.name
            tts.save(audio_path)
        
        # Read the generated audio file
        with open(audio_path, 'rb') as audio_file:
            audio_data = audio_file.read()
        
        # Clean up temporary file
        os.unlink(audio_path)
        
        # Convert to base64 for sending
        audio_base64 = base64.b64encode(audio_data).decode('utf-8')
        
        return jsonify({
            'status': 'success',
            'audio': audio_base64,
            'format': 'mp3'
        }), HTTPStatus.OK
        
    except Exception as e:
        logger.error(f"Error in text-to-speech: {str(e)}")
        return jsonify({'status': 'error', 'message': 'Error generating speech'}), HTTPStatus.INTERNAL_SERVER_ERROR

@app.errorhandler(404)
def not_found(error):
    return jsonify({
        'status': 'error',
        'message': 'Resource not found'
    }), HTTPStatus.NOT_FOUND

@app.errorhandler(405)
def method_not_allowed(error):
    return jsonify({
        'status': 'error',
        'message': 'Method not allowed'
    }), HTTPStatus.METHOD_NOT_ALLOWED

@app.errorhandler(500)
def internal_error(error):
    return jsonify({
        'status': 'error',
        'message': 'Internal server error'
    }), HTTPStatus.INTERNAL_SERVER_ERROR

if __name__ == "__main__":
    port = int(os.getenv('PORT', app.config['BACKEND_SERVER_PORT']))
    app.run(host='0.0.0.0', port=port)
